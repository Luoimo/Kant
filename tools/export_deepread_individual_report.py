from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "Individual_Report_DeepReadAgent.md"
OUTPUT = ROOT / "docs" / "Individual_Report_DeepReadAgent.docx"

ACCENT = RGBColor(31, 78, 121)
MUTED = RGBColor(90, 104, 120)
BORDER = "D9E2EC"
HEADER_FILL = "EEF3F8"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=130, bottom=90, end=130) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), BORDER)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(6)

    title = styles["Title"]
    title.font.name = "Arial"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = ACCENT
    title.paragraph_format.space_after = Pt(8)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Arial"
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = MUTED
    subtitle.paragraph_format.space_after = Pt(14)

    for name, size, before, after in [
        ("Heading 1", 16, 14, 6),
        ("Heading 2", 13, 10, 4),
        ("Heading 3", 11, 8, 3),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT if name != "Heading 3" else RGBColor(35, 35, 35)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    bullet = styles["List Bullet"]
    bullet.font.name = "Arial"
    bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    bullet.font.size = Pt(10.5)
    bullet.paragraph_format.left_indent = Inches(0.45)
    bullet.paragraph_format.first_line_indent = Inches(-0.2)
    bullet.paragraph_format.space_after = Pt(4)


def add_inline_markdown(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(60, 72, 88)
        else:
            paragraph.add_run(part)


def add_metadata_table(doc: Document, lines: list[str]) -> None:
    rows = []
    for line in lines:
        cleaned = line.strip()
        if not cleaned.startswith("**"):
            continue
        match = re.match(r"\*\*(.+?):\*\*\s*(.*)", cleaned)
        if match:
            rows.append((match.group(1), match.group(2)))

    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_borders(table)

    for idx, (label, value) in enumerate(rows):
        label_cell, value_cell = table.rows[idx].cells
        label_cell.width = Inches(2.0)
        value_cell.width = Inches(4.7)
        label_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        value_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(label_cell)
        set_cell_margins(value_cell)
        set_cell_shading(label_cell, HEADER_FILL)
        p = label_cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        r.bold = True
        r.font.color.rgb = ACCENT
        vp = value_cell.paragraphs[0]
        vp.paragraph_format.space_after = Pt(0)
        add_inline_markdown(vp, value)


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_table_separator(line: str) -> bool:
    stripped = line.strip()
    if not stripped.startswith("|"):
        return False
    parts = split_table_row(stripped)
    return all(re.fullmatch(r":?-{3,}:?", part or "") for part in parts)


def add_markdown_table(doc: Document, table_lines: list[str]) -> None:
    headers = split_table_row(table_lines[0])
    rows = [split_table_row(line) for line in table_lines[2:]]
    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_borders(table)
    widths_by_count = {
        3: [1.35, 3.0, 2.75],
        4: [1.35, 2.35, 1.85, 2.1],
    }
    widths = widths_by_count.get(col_count)

    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        if widths:
            cell.width = Inches(widths[idx])
        set_cell_shading(cell, HEADER_FILL)
        set_cell_margins(cell, top=100, bottom=100, start=120, end=120)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_inline_markdown(p, f"**{text}**")
        for run in p.runs:
            run.font.size = Pt(8.5)

    for r_idx, row in enumerate(rows, 1):
        for c_idx in range(col_count):
            cell = table.rows[r_idx].cells[c_idx]
            if widths:
                cell.width = Inches(widths[c_idx])
            set_cell_margins(cell, top=90, bottom=90, start=120, end=120)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_inline_markdown(p, row[c_idx] if c_idx < len(row) else "")
            for run in p.runs:
                run.font.size = Pt(8.2)
                if run.font.name == "Courier New":
                    run.font.size = Pt(7.4)

    doc.add_paragraph()


def build_docx() -> None:
    raw = SOURCE.read_text(encoding="utf-8")
    lines = raw.splitlines()

    doc = Document()
    style_document(doc)

    header = doc.sections[0].header
    hp = header.paragraphs[0]
    hp.text = "Kant Individual Project Report"
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.runs[0].font.name = "Arial"
    hp.runs[0].font.size = Pt(9)
    hp.runs[0].font.color.rgb = MUTED

    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    add_page_number(fp)
    for run in fp.runs:
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED

    title_text = lines[0].lstrip("# ").strip()
    title_para = doc.add_paragraph(style="Title")
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_para.add_run(title_text)

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("DeepReadAgent design, implementation, validation, AI safety, and responsible AI reflection")

    metadata_lines = []
    body_start = 1
    for i, line in enumerate(lines[1:], 1):
        if line.startswith("## "):
            body_start = i
            break
        if line.strip():
            metadata_lines.append(line)
    add_metadata_table(doc, metadata_lines)
    doc.add_paragraph()

    i = body_start
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            table_lines = [stripped, lines[i + 1].strip()]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            add_markdown_table(doc, table_lines)
            continue
        elif stripped.startswith("## "):
            doc.add_paragraph(stripped[3:].strip(), style="Heading 1")
        elif stripped.startswith("### "):
            doc.add_paragraph(stripped[4:].strip(), style="Heading 2")
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_markdown(p, stripped[2:].strip())
        elif re.match(r"\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            add_inline_markdown(p, re.sub(r"^\d+\.\s+", "", stripped))
        else:
            p = doc.add_paragraph()
            add_inline_markdown(p, stripped)
        i += 1

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_docx()
    print(OUTPUT)
