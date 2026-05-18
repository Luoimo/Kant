from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/Team22_Phase0_1_First_Presentation_Report.docx")


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_in: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="DADCE0") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)

    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        cell = hdr[i]
        set_cell_fill(cell, LIGHT_FILL)
        set_cell_width(cell, widths[i])
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(31, 61, 97)

    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cell = cells[i]
            set_cell_width(cell, widths[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            run.font.size = Pt(9)
    doc.add_paragraph()
    return table


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="C9D6E2")
    cell = table.rows[0].cells[0]
    set_cell_fill(cell, CALLOUT_FILL)
    set_cell_width(cell, 6.5)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    r.font.size = Pt(10.5)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    r2.font.size = Pt(10)
    doc.add_paragraph()


def add_code_block(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="D0D7DE")
    cell = table.rows[0].cells[0]
    set_cell_fill(cell, "F6F8FA")
    set_cell_width(cell, 6.5)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(8.5)
    doc.add_paragraph()


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(8)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = MUTED
    subtitle.paragraph_format.space_after = Pt(12)

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.color.rgb = BLUE
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.color.rgb = BLUE
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)

    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.color.rgb = DARK_BLUE
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "Kant Phase 0+1 First Presentation Report"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Page ")
    run.font.size = Pt(9)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)


def build_doc() -> None:
    doc = Document()
    setup_styles(doc)
    add_header_footer(doc)

    doc.add_paragraph("Kant", style="Title")
    doc.add_paragraph("Phase 0+1 First Presentation Report", style="Subtitle")
    meta = doc.add_paragraph()
    meta.add_run("Project: ").bold = True
    meta.add_run("Agentic AI Deep Reading and Knowledge Management System\n")
    meta.add_run("Team: ").bold = True
    meta.add_run("Team 22\n")
    meta.add_run("Team Members: ").bold = True
    meta.add_run("GAO MINGZE, HUANG YINUO, REN ZHENG, WANG WENZHE, SHAO QI\n")
    meta.add_run("Submission Context: ").bold = True
    meta.add_run("First Project Presentation / Post Phase 0+1 submission draft")

    add_callout(
        doc,
        "Executive positioning",
        "This report is written for the first project presentation. It follows the NUS-ISS presentation guideline sequence and documents the current Phase 0+1 evidence: scope, backlog, sprint progress capture, architecture, design, implementation, CI/CD, testing, security, risks, and next-phase work.",
    )

    doc.add_heading("1. Project Overview and Introduction", level=1)
    doc.add_paragraph(
        "Kant is an AI-assisted deep reading and knowledge management system for EPUB books. The system allows users to upload books, read them in a browser-based reader, ask contextual questions about the current book or selected passage, and receive evidence-grounded answers with citations. It also supports follow-up question generation, note creation, conversation history, user isolation, administration, and audit logging."
    )
    doc.add_paragraph(
        "The project addresses three practical problems in long-form reading: readers struggle to locate evidence inside long books, generic chatbots can hallucinate when asked about specific text, and useful reading insights are often not captured as reusable notes. Kant solves these problems through a software engineering system that combines EPUB ingestion, hybrid retrieval, graph-based knowledge support, controlled AI orchestration, and testable security controls."
    )
    add_callout(
        doc,
        "Current Phase 0+1 status",
        "The current MVP has implemented the core reading workflow, book ingestion, authenticated user access, conversation management, RAG-based answering, citations, multi-agent orchestration, admin/audit APIs, automated backend tests, AI security tests, and GitHub Actions CI/CD. Remaining Phase 2 work focuses on production hardening, monitoring, E2E/stress testing, evaluation, and final presentation/report refinement.",
    )

    doc.add_heading("2. Overall Scope and Use Cases", level=1)
    doc.add_paragraph(
        "The system scope covers the end-to-end reading assistant workflow rather than only an isolated AI function. The main actors are the reader/member, the administrator, the LLM provider, Notion, and the data stores used for metadata, vectors, graph relationships, and files."
    )
    add_table(
        doc,
        ["Actor", "Primary Use Cases", "Current Phase 0+1 Status"],
        [
            ["Reader / Member", "Login, browse library, upload EPUB, open reader, select text, ask book questions, view citations, view follow-up questions, manage conversations.", "Core workflow implemented."],
            ["Administrator", "List users, inspect user books/conversations, review audit trail.", "Admin APIs and audit logging implemented; UI hardening remains Phase 2."],
            ["LLM Provider", "Generate answers, perform routing, review answer quality, extract note structure, generate follow-up questions.", "Integrated through OpenAI-compatible client."],
            ["Notion API", "Persist and organise reading notes generated from Q&A.", "Implemented as NoteAgent backend/tool target."],
            ["Storage Services", "Store books, covers, catalog metadata, embeddings, graph entities, conversations, and audit logs.", "PostgreSQL/ChromaDB/Neo4j/OSS-or-local patterns implemented or designed."],
        ],
        [1.25, 3.75, 1.5],
    )
    add_code_block(
        doc,
        """Use case overview

Reader/Member:
  - Authenticate
  - Upload EPUB
  - Browse library
  - Read EPUB and track current chapter/selection
  - Ask AI question about book or selected text
  - Receive streamed answer, citations, and follow-up questions
  - Continue or clear conversation history

Admin:
  - View users
  - View user books
  - View user conversations
  - Generate audit trail for governance

External systems:
  - LLM API for generation and review
  - Notion API for note persistence
  - ChromaDB / Neo4j / PostgreSQL / OSS for retrieval and storage"""
    )

    doc.add_heading("3. Project Roadmap and Key Milestones", level=1)
    add_table(
        doc,
        ["Phase", "Target Outcomes", "Evidence / Deliverables"],
        [
            ["Phase 0: Product and Architecture Planning", "Define business objective, system scope, initial architecture, backlog, key risks, and technical approach.", "Project requirement alignment, architecture direction, module boundaries, backlog items."],
            ["Phase 1: MVP Construction", "Implement a working slice covering EPUB ingestion, reading UI, authenticated chat workflow, retrieval, citations, AI orchestration, testing, and CI/CD.", "Working frontend/backend, tests, CI workflow, security tests, evidence report."],
            ["Phase 2: Production Hardening and Finalisation", "Strengthen deployment, monitoring, evaluation, stress/E2E tests, report quality, demo scripts, and final submission artifacts.", "Final report, demo videos, refined diagrams, production hardening evidence."],
        ],
        [1.4, 3.1, 2.0],
    )

    doc.add_heading("4. Complete Project Backlog", level=1)
    add_table(
        doc,
        ["Backlog Area", "Representative Items", "Status"],
        [
            ["Frontend Reading Experience", "Library page, EPUB reader, reading progress, selected text context, chat side panel.", "Done"],
            ["Book Ingestion", "EPUB upload, temporary processing, cover extraction, chunking, source metadata, catalog registration.", "Done"],
            ["RAG and Knowledge Retrieval", "ChromaDB vector search, BM25-style keyword search, hybrid retrieval, reranking, citation construction, optional graph retrieval.", "Done / continuing quality tuning"],
            ["AI Orchestration", "RouterAgent, DeepReadAgent, CriticAgent, NoteAgent, FollowupAgent, SSE streaming, memory context.", "Done for MVP"],
            ["Identity and Access", "Login, JWT access, refresh support, member/admin role checks, user-owned books and conversations.", "Done"],
            ["Storage and Data Management", "PostgreSQL catalog schema, conversation catalog, checkpoint store, OSS/local fallback, audit logs.", "Done / deployment hardening planned"],
            ["Admin and Governance", "Admin APIs for users/books/conversations and audit log creation.", "Implemented; UI/demo polish planned"],
            ["CI/CD and Testing", "GitHub Actions lint, backend tests, AI security tests, frontend build, test artifacts.", "Done"],
            ["Security Hardening", "Lakera Guard integration, local fallback rules, prompt injection tests, secret/filesystem/code-execution tests.", "Done / red-team expansion planned"],
            ["Phase 2 Evaluation", "RAGAS or equivalent retrieval/answer evaluation, E2E testing, stress testing, monitoring dashboard.", "Planned"],
        ],
        [1.45, 4.1, 0.95],
    )

    doc.add_heading("5. Sprint Effort to Date", level=1)
    doc.add_paragraph(
        "The repository contains implementation and evidence artifacts, but it does not contain authoritative per-member effort logs. The table below should be completed with the team’s actual effort figures before submission. The work categories reflect the implemented modules and can be used as the sprint-effort capture structure required by the presentation guideline."
    )
    add_table(
        doc,
        ["Sprint / Period", "Planned Work", "Completed Work", "Owner(s)", "Estimated Hours", "Actual Hours"],
        [
            ["Sprint 0", "Requirements, proposal alignment, architecture direction, backlog definition.", "Project scope and architecture direction established.", "Team", "TBC", "TBC"],
            ["Sprint 1", "EPUB ingestion, library/reader UI, RAG foundation.", "Book upload, reader view, ChromaDB retrieval, chunking and cleaning tests.", "Team", "TBC", "TBC"],
            ["Sprint 2", "Chat workflow, agents, citations, streaming.", "DeepReadAgent workflow, SSE response, citations, follow-up generation, note hook.", "Team", "TBC", "TBC"],
            ["Sprint 3", "Security, auth, storage, CI/CD.", "User isolation, PostgreSQL catalog, admin/audit APIs, AI security tests, GitHub Actions.", "Team", "TBC", "TBC"],
            ["Sprint 4 / Phase 2", "Deployment, monitoring, E2E/stress tests, final report and demo.", "Planned next-phase work.", "Team", "TBC", "TBC"],
        ],
        [1.0, 1.7, 1.7, 0.8, 0.65, 0.65],
    )

    doc.add_heading("6. Technology Stack", level=1)
    add_table(
        doc,
        ["Layer", "Technology", "Purpose"],
        [
            ["Frontend", "Vue 3, Vite, Naive UI", "Library, EPUB reader, chat panel, progress and interaction UI."],
            ["Backend API", "FastAPI, Python 3.12", "REST/SSE APIs, authentication, book ingestion, chat orchestration, admin APIs."],
            ["AI / Agent Runtime", "LangGraph, LangChain, OpenAI-compatible API", "Tool-using agent workflow, routing, answer generation, review, notes, follow-up questions."],
            ["Retrieval", "ChromaDB, hybrid retriever, BM25-style keyword retrieval, reranker", "Evidence retrieval and citation grounding."],
            ["Knowledge Graph", "Neo4j, HanLP/LLM extraction path", "Concept/entity relationship support for deeper reading analysis."],
            ["Databases and Files", "PostgreSQL, checkpoint store, OSS/local files", "Users, books, conversations, audit logs, EPUBs, covers, chat state."],
            ["Security and Governance", "Lakera Guard, local safety fallback, JWT, role-based access, audit logs", "AI safety, access control, traceability."],
            ["DevOps", "GitHub Actions, Pytest, coverage, Dockerfile", "Automated testing, build checks, artifact evidence, deployable packaging."],
        ],
        [1.3, 1.8, 3.4],
    )

    doc.add_heading("7. Architectural Constraints and Decisions", level=1)
    add_table(
        doc,
        ["Constraint / Risk", "Decision", "Expected Benefit"],
        [
            ["Answers must be verifiable.", "Use retrieval-grounded generation and return citations from source chunks.", "Improves explainability and user trust."],
            ["Multiple users must not access each other’s books or conversations.", "Use authenticated APIs and owner_user_id checks across catalog and conversation access.", "Supports privacy and governance."],
            ["LLMs may hallucinate or overgeneralise.", "Use RAG, citations, CriticAgent review, and conservative prompts.", "Reduces unsupported answers and surfaces risk."],
            ["External safety services may be unavailable.", "Use Lakera Guard when configured and local rule-based fallback otherwise.", "Keeps safety checks active during failures."],
            ["EPUB files and covers may need cloud storage.", "Support OSS URI storage with local fallback.", "Keeps local development simple while preserving cloud path."],
            ["Distributed microservices may add operational complexity too early.", "Use a modular monolith for Phase 0+1 and define clear domain boundaries.", "Keeps MVP cohesive while allowing future extraction."],
        ],
        [1.7, 2.5, 2.3],
    )

    doc.add_heading("8. Logical Architecture", level=1)
    doc.add_paragraph(
        "Kant uses a layered logical architecture. The frontend owns the reading interaction and passes contextual information such as book ID, selected text, current chapter, locale, and conversation ID to the backend. The backend validates user access, performs safety checks, orchestrates the AI workflow, retrieves evidence, and writes results to storage or downstream tools."
    )
    add_code_block(
        doc,
        """Logical architecture

User Interface Layer
  Vue LibraryView / ReaderView / ReaderChat

API and Access Layer
  FastAPI routers: auth, books, conversations, chat, admin
  JWT authentication and member/admin dependencies

Security and Governance Layer
  Lakera Guard + local fallback input filter
  role checks, owner checks, audit logging

Domain and AI Orchestration Layer
  Book ingestion, conversation management
  RouterAgent -> DeepReadAgent -> CriticAgent / NoteAgent / FollowupAgent

Retrieval and Knowledge Layer
  EPUB extractor, cleaner, chunker
  ChromaDB, hybrid retriever, citations, Neo4j graph

Storage and Integration Layer
  PostgreSQL catalog, checkpoint store, OSS/local files, Notion API, LLM API"""
    )

    doc.add_heading("9. Physical and Deployment Architecture", level=1)
    doc.add_paragraph(
        "The current implementation is runnable as a local or containerised web application, with FastAPI serving the backend and Vue/Vite serving the frontend. The intended production architecture separates compute from persistent services: the frontend and backend can run in containers, while metadata, vectors, graph data, files, model APIs, and observability services are externalised."
    )
    add_code_block(
        doc,
        """Physical deployment view

Browser
  -> Vue frontend
  -> FastAPI backend container
       -> PostgreSQL: users, books, conversations, notes, audit_logs
       -> ChromaDB: embeddings and book chunks
       -> Neo4j: book entities and concept graph
       -> OSS/local storage: EPUB files and covers
       -> OpenAI-compatible LLM API: generation and review
       -> Notion API: structured note persistence
       -> GitHub Actions: CI checks and artifacts"""
    )
    doc.add_paragraph(
        "For the first presentation, the honest deployment position is that Phase 0+1 focuses on a working MVP plus CI evidence. Full production deployment, monitoring dashboards, autoscaling, and stress testing are planned Phase 2 hardening items."
    )

    doc.add_heading("10. Domain Design and Module Boundaries", level=1)
    doc.add_paragraph(
        "The current architecture is best described as a modular monolith rather than a distributed microservice system. This is intentional for the MVP because the user request path needs tight orchestration among API validation, safety, retrieval, LLM calls, streaming, and post-processing. The codebase still follows domain-oriented boundaries so that modules can be tested and evolved independently."
    )
    add_table(
        doc,
        ["Domain Module", "Responsibility", "Representative Files"],
        [
            ["Identity and Access", "Login, token handling, role enforcement, user isolation.", "backend/api/auth.py, backend/auth/*, backend/api/deps.py"],
            ["Library Management", "Book upload, listing, file access, deletion, cover handling.", "backend/api/books.py, backend/storage/book_catalog.py"],
            ["Reading Session and Conversation", "Reader context, conversation creation/history, chat state.", "frontend/src/views/ReaderView.vue, backend/api/conversations.py, backend/storage/checkpoint_store.py"],
            ["AI Orchestration", "Route, answer, review, note, and follow-up workflows.", "backend/agents/*, backend/api/chat.py"],
            ["Retrieval and Knowledge", "EPUB extraction, chunking, vector search, graph retrieval, citations.", "backend/rag/*, backend/graph/*, backend/xai/citation.py"],
            ["Admin and Audit", "Admin views and audit logging for governance.", "backend/api/admin.py, backend/storage/audit_log_catalog.py"],
        ],
        [1.5, 2.4, 2.6],
    )

    doc.add_heading("11. Software Design", level=1)
    doc.add_heading("11.1 Book Ingestion Sequence", level=2)
    add_code_block(
        doc,
        """Book ingestion sequence

User -> Frontend: upload EPUB
Frontend -> FastAPI /books/upload: multipart file + token
FastAPI -> Auth dependency: validate member
FastAPI -> temp file: save upload stream
FastAPI -> ChromaStore: extract, clean, chunk, embed, store vectors
FastAPI -> OSS/local file storage: persist EPUB and cover
FastAPI -> BookCatalog/PostgreSQL: register metadata
FastAPI -> Neo4j: upsert book and graph relationships
FastAPI -> Frontend: return book_id, source, collection, chunk counts"""
    )
    doc.add_heading("11.2 Contextual Reading Q&A Sequence", level=2)
    add_code_block(
        doc,
        """Contextual Q&A sequence

User -> ReaderChat: ask question with optional selected text/current chapter
ReaderChat -> /api/user/chat/stream: query, book_id, conversation_id, context
FastAPI -> Safety Filter: Lakera Guard or local fallback
FastAPI -> ConversationCatalog: verify conversation belongs to user and book
FastAPI -> Mem0/checkpoint store: retrieve relevant memory and short-term state
FastAPI -> RouterAgent: identify intent and optimised query
FastAPI -> DeepReadAgent: build dynamic context and call retrieval tools
DeepReadAgent -> HybridRetriever: retrieve book evidence from ChromaDB/BM25/graph support
DeepReadAgent -> LLM: generate answer grounded in evidence
Citation Builder -> Frontend: return citations and document snippets
CriticAgent -> Frontend: stream review notes when needed
NoteAgent -> Notion: persist structured reading note
FollowupAgent -> Frontend: return reflective follow-up questions"""
    )

    doc.add_heading("12. Database, NoSQL, and Collection Design", level=1)
    add_table(
        doc,
        ["Store", "Objects / Collections", "Purpose"],
        [
            ["PostgreSQL", "users, books, notes, conversations, audit_logs", "Authoritative metadata, identity, ownership, governance, and conversation catalog."],
            ["ChromaDB", "book chunk embeddings with metadata such as book_id, owner_user_id, source, title, section", "Vector retrieval for evidence-grounded answers."],
            ["Neo4j", "Book, concept/entity, chapter, relationship paths", "Graph-based expansion and concept relationship support."],
            ["Checkpoint Store", "conversation thread state and AI messages", "Short-term multi-turn chat state."],
            ["OSS / Local Files", "EPUB files and cover images", "Binary file storage with cloud-ready signed URL support."],
            ["Notion", "Generated reading notes and cross-book knowledge entries", "User-facing knowledge management destination."],
        ],
        [1.3, 2.8, 2.4],
    )

    doc.add_heading("13. CI/CD Pipeline", level=1)
    doc.add_paragraph(
        "The GitHub Actions workflow provides a Phase 0+1 LLMSecOps foundation. It separates syntax/lint preparation, backend core tests, AI security tests, frontend build, artifact upload, and a final success gate."
    )
    add_code_block(
        doc,
        """CI/CD pipeline

Push / Pull Request
  -> Preparation & Lint
  -> Backend Core Tests
       pytest tests/ --ignore=tests/security --cov=. --cov-report=xml --junitxml=core-results.xml
  -> Backend AI Security Tests
       pytest tests/security --junitxml=security-results.xml
  -> Frontend Build
       npm install / npm ci, npm run build
  -> Upload Artifacts
       core-results.xml, coverage.xml, security-results.xml
  -> CI Pipeline Passed gate"""
    )
    add_callout(
        doc,
        "Phase 2 CI hardening note",
        "The project development rule states that local runs should use .venv312. The current CI workflow uses Python 3.10. Before final submission, the CI Python version should be aligned with Python 3.12 to remove environment drift.",
    )

    doc.add_heading("14. Live Demo Plan", level=1)
    add_numbered(
        doc,
        [
            "Log in as a member user.",
            "Open the library page and select an uploaded EPUB.",
            "Show the split reader view with the EPUB reader on the left and AI chat on the right.",
            "Select a text passage and ask a contextual question.",
            "Show the streamed answer and tool/status updates.",
            "Show citations returned with the answer.",
            "Show follow-up questions generated by the system.",
            "Run a prompt-injection or secret-leakage example and show that the request is rejected.",
            "Show the CI/test evidence report with backend and security test results.",
        ],
    )

    doc.add_heading("15. Testing Summary", level=1)
    doc.add_paragraph(
        "The current evidence report records 161 passing backend tests and 56.5% line coverage. The tests validate core backend behaviour, RAG components, storage, API paths, authentication, and AI security controls. The coverage figure is adequate for Phase 0+1 evidence but should be improved for final submission, especially around frontend E2E, stress testing, and production monitoring paths."
    )
    add_table(
        doc,
        ["Test Type", "Current Evidence", "Phase 2 Improvement"],
        [
            ["Unit Testing", "RAG chunker, cleaner, Chroma store, retriever, note agent, storage components.", "Increase coverage around edge cases and graph extraction."],
            ["Integration Testing", "Chat API, conversation API, admin API, auth API, storage integration tests.", "Add fuller user journey integration tests with realistic data."],
            ["AI Security Testing", "Prompt injection, secrets leakage, filesystem access, command execution, off-topic handling.", "Expand red-team cases and add regression samples from demo feedback."],
            ["End-to-End Testing", "Manual demo path currently available.", "Automate login -> library -> reader -> chat -> citation flow."],
            ["Stress Testing", "Not yet complete in Phase 0+1.", "Add load tests for streaming chat and retrieval-heavy workloads."],
        ],
        [1.35, 3.2, 1.95],
    )

    doc.add_heading("16. Management Concerns, Issues, and Mitigations", level=1)
    add_table(
        doc,
        ["Concern", "Impact", "Mitigation"],
        [
            ["Large system scope", "Risk of spreading effort across too many features.", "Use MVP-first scope: reading, Q&A, citations, safety, tests. Keep monitoring and stress testing for Phase 2."],
            ["External API dependencies", "Live demo may fail if LLM, Notion, Neo4j, or OSS credentials are unavailable.", "Prepare fixed demo data, local fallback path, and recorded demo videos."],
            ["Team effort reporting", "Presentation requires estimated vs actual effort, but raw logs must be consolidated.", "Collect per-member hours and update the effort table before submission."],
            ["Integration risk", "Frontend, backend, retrieval, and storage changes may break each other.", "Keep CI mandatory and add E2E tests in Phase 2."],
            ["Report and evidence alignment", "Claims may be challenged if not supported by artifacts.", "Map each claim to code, tests, CI, screenshots, or demo video evidence."],
        ],
        [1.7, 2.25, 2.55],
    )

    doc.add_heading("17. Technical Concerns, Issues, and Mitigations", level=1)
    add_table(
        doc,
        ["Technical Concern", "Risk", "Mitigation / Next Action"],
        [
            ["Retrieval quality", "Weak retrieval can produce weak or unsupported answers.", "Tune chunking, hybrid retrieval, reranking, and add RAG evaluation."],
            ["Graph extraction accuracy", "Incorrect entity relationships may mislead graph-based context.", "Validate graph extraction on selected books and expose graph as supporting, not sole, evidence."],
            ["LLM latency and cost", "Streaming and multiple agents can increase latency and token usage.", "Use asynchronous post-processing and measure latency/token metrics in Phase 2."],
            ["Storage portability", "Local and cloud paths may diverge.", "Keep OSS/local fallback explicit and test both modes."],
            ["Environment drift", "Local .venv312 differs from CI Python 3.10.", "Align CI to Python 3.12."],
        ],
        [1.6, 2.2, 2.7],
    )

    doc.add_heading("18. Security Concerns, Mitigations, and Solutions", level=1)
    add_table(
        doc,
        ["Security Risk", "Implemented Control", "Evidence"],
        [
            ["Prompt injection", "Lakera Guard integration with local keyword fallback.", "backend/security/input_filter.py and backend/tests/security/test_security.py"],
            ["Secrets leakage", "Detection of API-key/token patterns before agent execution.", "Security test cases for OpenAI-style and GitHub token patterns."],
            ["Filesystem or command abuse", "Reject filesystem access and command execution requests; agents receive scoped tools only.", "Security tests for /etc/passwd, rm -rf, os.system, cmd /c."],
            ["Hallucination", "RAG grounding, citations, CriticAgent review, evidence-aware prompts.", "Citation builder, DeepReadAgent retrieval, CriticAgent workflow."],
            ["Cross-user data leakage", "Authenticated APIs, owner_user_id checks, conversation-book ownership validation.", "Auth and conversation API tests; backend/api/chat.py ownership checks."],
            ["Admin misuse or opaque operations", "Role-based admin dependency and audit log writes.", "backend/api/admin.py and audit_logs schema."],
        ],
        [1.55, 2.55, 2.4],
    )

    doc.add_heading("19. AOB and Next Steps", level=1)
    add_bullets(
        doc,
        [
            "Complete the first presentation slide deck using this report as the content source.",
            "Replace the sprint effort table with actual member effort logs before submission.",
            "Prepare four short videos if required: app demo, CI/CD demo, testing evidence demo, and security demo.",
            "Align CI Python version with the project .venv312 development rule.",
            "Add E2E and stress testing evidence for final submission.",
            "Refine logical, physical, deployment, database, and sequence diagrams into presentation-ready visuals.",
            "Incorporate lecturer feedback from the first presentation into the Phase 0+1 report and final project report.",
        ],
    )

    doc.add_heading("Appendix A. Evidence Mapping", level=1)
    add_table(
        doc,
        ["Requirement / Guideline Area", "Project Evidence"],
        [
            ["Project overview and scope", "README, current frontend/backend workflow, use cases in this report."],
            ["Roadmap, backlog, sprint effort", "Backlog/status table and effort capture template."],
            ["Logical and physical architecture", "Layered architecture and deployment component diagrams in this report."],
            ["Software design", "Book ingestion and Q&A sequence designs."],
            ["DB/NoSQL design", "PostgreSQL schema, ChromaDB, Neo4j, checkpoint store, OSS/local files."],
            ["CI/CD", ".github/workflows/ci.yml."],
            ["Testing", "docs/test_evidence_report.html showing 161/161 passing tests and 56.5% line coverage."],
            ["Security", "backend/security/input_filter.py and backend/tests/security/test_security.py."],
            ["Agentic AI requirement", "RouterAgent, DeepReadAgent, CriticAgent, NoteAgent, FollowupAgent integrated in backend/api/chat.py."],
            ["Explainable and Responsible AI", "Citation builder, retrieved evidence, CriticAgent, audit logging and access control."],
        ],
        [2.25, 4.25],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
