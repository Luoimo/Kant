from docx import Document


DOCX_PATH = "docs/Kant_System_Architecture_Sections.docx"


replacement_overview = {
    7: (
        "For production deployment, Kant is designed as a containerised modular system on Alibaba Cloud. "
        "The system can be deployed on Alibaba Cloud Container Service for Kubernetes (ACK), with container "
        "images built through CI/CD and stored in Alibaba Cloud Container Registry (ACR). The key deployment "
        "unit is not a single-agent pod. Instead, each backend application pod contains the complete five-agent "
        "runtime pipeline, so one pod can process an end-to-end reading request independently."
    ),
    8: (
        "This architecture supports the project’s key requirements: multi-agent collaboration, explainability, "
        "AI security, scalable deployment, monitoring, and LLMSecOps. It also avoids excessive network hops "
        "between agents. The system scales by increasing the number of identical pipeline pod replicas, while "
        "shared dependencies such as databases, vector storage, graph storage, queues, and observability services "
        "scale independently."
    ),
}


physical_paragraphs = {
    20: (
        "The production deployment of Kant is designed as a replicated pipeline architecture on Alibaba Cloud. "
        "The system runs on Alibaba Cloud Container Service for Kubernetes (ACK), while container images are "
        "built through CI/CD and stored in Alibaba Cloud Container Registry (ACR). ACK provides container "
        "orchestration, service discovery, rolling updates, horizontal scaling, workload isolation, and integration "
        "with Alibaba Cloud networking, storage, monitoring, and security services."
    ),
    21: (
        "External traffic first reaches Alibaba Cloud CDN, WAF, and Anti-DDoS services. CDN accelerates static "
        "frontend assets, while WAF and Anti-DDoS protect the system from common web attacks and volumetric "
        "attacks. Requests are then routed through API Gateway, which provides API authentication, rate limiting, "
        "request validation, and traffic control. ALB/SLB and ACK Ingress route traffic into the Kubernetes cluster."
    ),
    22: (
        "Inside ACK, Kant uses a small number of service deployments, but the core AI capability is packaged as "
        "a full pipeline pod. The frontend-web deployment serves the Vue application. The main backend deployment, "
        "kant-pipeline, exposes the FastAPI endpoints and contains the Security Gateway, RouterAgent, DeepReadAgent, "
        "CriticAgent, NoteAgent, FollowupAgent, orchestration logic, retrieval adapters, citation generation, and "
        "streaming response handler in the same pod. Therefore, Kubernetes scales complete agentic pipeline replicas "
        "rather than scaling one pod per agent."
    ),
    23: (
        "Within each kant-pipeline pod, the request is handled in a local sequence. The Security Gateway screens "
        "the request for prompt injection, adversarial input, and policy risks. RouterAgent classifies the intent "
        "and prepares the query. DeepReadAgent calls shared retrieval services and model endpoints to generate an "
        "evidence-grounded answer. CriticAgent checks citation grounding, hallucination risk, and answer quality. "
        "NoteAgent writes structured notes when requested, and FollowupAgent proposes follow-up questions or reading "
        "paths. Agent communication is primarily in-process, while durable state and asynchronous tasks are written "
        "to Redis, RDS/PolarDB, OSS, NAS, Neo4j, and the vector database."
    ),
    24: (
        "Persistent data is separated from compute containers. OSS stores raw EPUB files, covers, exported reports, "
        "test evidence, and security scan artifacts. RDS MySQL or PolarDB stores user records, book metadata, chunk "
        "metadata, chat indexes, task status, and audit logs. Redis stores sessions, short-term cache, retrieval "
        "cache, distributed locks, and asynchronous task queues. NAS stores the Obsidian-style Markdown knowledge "
        "vault. Neo4j stores concept graphs and relationships between books, chapters, concepts, and notes. The "
        "vector database stores document embeddings and supports semantic retrieval."
    ),
    25: (
        "Secrets and permissions are managed through RAM, KMS, and Kubernetes Secrets. The frontend does not contain "
        "API keys. Only kant-pipeline pods and controlled worker pods can access OpenAI, Lakera Guard, LangSmith, "
        "database credentials, and storage credentials. NetworkPolicy is used to restrict lateral communication "
        "between pods, so data services are accessible only to approved backend workloads."
    ),
    26: (
        "Observability is implemented through multiple layers. SLS collects application, security, and audit logs. "
        "ARMS provides application performance monitoring and distributed tracing. CloudMonitor tracks infrastructure "
        "metrics. Prometheus and Grafana provide Kubernetes and custom business metrics. LangSmith records LLM traces, "
        "latency, token usage, and cost. Each pipeline pod emits per-agent spans and events, so the team can inspect "
        "RouterAgent, DeepReadAgent, CriticAgent, NoteAgent, and FollowupAgent behaviour even though they run inside "
        "the same pod."
    ),
    27: (
        "Autoscaling is based on complete pipeline replicas. HPA scales frontend-web and kant-pipeline pods based on "
        "CPU, memory, QPS, request queue depth, token throughput, and p95 latency. Asynchronous ingestion or batch "
        "jobs can be handled by separate worker pods through Redis Stream, RocketMQ, or KEDA-driven queues, but the "
        "online reading and question-answering path remains a full five-agent pipeline inside every kant-pipeline pod. "
        "GPU workloads, if self-hosted, run in a dedicated GPU ECS node pool and are accessed as shared inference "
        "endpoints by all pipeline pods."
    ),
    28: (
        "Deployment architecture:\n"
        "User Browser\n"
        "   |\n"
        "   v\n"
        "Alibaba Cloud CDN\n"
        "   |\n"
        "   v\n"
        "WAF + Anti-DDoS\n"
        "   |\n"
        "   v\n"
        "API Gateway\n"
        "   |\n"
        "   v\n"
        "ALB / SLB\n"
        "   |\n"
        "   v\n"
        "ACK Ingress Controller\n"
        "   |\n"
        "   +--> frontend-web pods\n"
        "   |\n"
        "   +--> kant-pipeline pods  x N replicas\n"
        "            |\n"
        "            +--> Security Gateway\n"
        "            +--> RouterAgent\n"
        "            +--> DeepReadAgent\n"
        "            +--> CriticAgent\n"
        "            +--> NoteAgent\n"
        "            +--> FollowupAgent\n"
        "            |\n"
        "            +--> Retrieval adapters: Vector DB / BM25 Index / Neo4j / Redis Cache\n"
        "            +--> Model endpoints: OpenAI API / Self-hosted LLM or Embedding GPU Service\n"
        "            +--> State and storage: RDS or PolarDB / Redis / OSS / NAS\n"
        "            +--> Observability: SLS / ARMS / CloudMonitor / Prometheus / Grafana / LangSmith\n"
        "   |\n"
        "   +--> optional ingestion-worker pods\n"
        "            +--> EPUB extraction / cleaning / chunking / embedding / indexing"
    ),
}


table_rows = [
    ("Component", "Alibaba Cloud Service", "Responsibility"),
    ("Container orchestration", "ACK", "Runs and scales replicated frontend, pipeline, and worker pods"),
    ("Image registry", "ACR", "Stores versioned frontend, kant-pipeline, and worker container images"),
    ("Static asset delivery", "CDN", "Accelerates frontend assets"),
    ("Edge protection", "WAF, Anti-DDoS", "Protects against web and DDoS attacks"),
    ("API entry", "API Gateway, ALB/SLB, Ingress", "Authenticates, rate-limits, validates, and routes traffic"),
    ("Full agent pipeline pod", "ACK Deployment: kant-pipeline", "Contains Security Gateway plus Router, DeepRead, Critic, Note, and Followup agents in every pod replica"),
    ("Asynchronous workers", "ACK worker Deployment, KEDA, Redis Stream / RocketMQ", "Handles ingestion, indexing, note export, and other non-blocking jobs"),
    ("Object storage", "OSS", "Stores EPUBs, covers, reports, and security evidence"),
    ("Metadata database", "RDS MySQL / PolarDB", "Stores users, books, chunks, tasks, chat indexes, and audit records"),
    ("Cache, locks, and queue", "Redis", "Stores sessions, retrieval cache, distributed locks, and async queues"),
    ("Shared notes", "NAS", "Stores the Markdown knowledge vault"),
    ("Graph database", "Neo4j on ACK or ECS", "Stores concept relationships"),
    ("Vector database", "Milvus / AnalyticDB Vector / vector DB", "Stores embeddings and supports semantic retrieval"),
    ("Secrets", "KMS, RAM, Kubernetes Secrets", "Manages credentials and least-privilege access"),
    ("Logs", "SLS", "Collects application, agent, security, and audit logs"),
    ("APM", "ARMS", "Tracks latency and service performance"),
    ("Metrics", "CloudMonitor, Prometheus, Grafana", "Monitors infrastructure, pod replicas, queues, and custom metrics"),
    ("LLM observability", "LangSmith", "Tracks LLM traces, per-agent spans, token usage, and cost"),
    ("GPU inference monitoring", "DCGM Exporter, Prometheus", "Monitors optional self-hosted model, embedding, and reranker endpoints"),
]


def main() -> None:
    doc = Document(DOCX_PATH)

    for idx, text in replacement_overview.items():
        doc.paragraphs[idx].text = text

    for idx, text in physical_paragraphs.items():
        doc.paragraphs[idx].text = text

    table = doc.tables[0]
    while len(table.rows) < len(table_rows):
        table.add_row()
    for row_idx, row_values in enumerate(table_rows):
        for col_idx, value in enumerate(row_values):
            table.rows[row_idx].cells[col_idx].text = value

    for row_idx in range(len(table.rows) - 1, len(table_rows) - 1, -1):
        tr = table.rows[row_idx]._tr
        tr.getparent().remove(tr)

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    main()
